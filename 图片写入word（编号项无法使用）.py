from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
from openpyxl import load_workbook
from PIL import Image


def set_global_font(doc, font_name='宋体', font_size=Pt(10.5)):
    """设置全局默认字体"""
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    doc.styles['Normal'].font.size = font_size


def extract_serial_number(filename):
    """从文件名中提取序号"""
    match = re.search(r'pt(\d+)', filename)
    if match:
        return match.group(1)
    else:
        return filename


def add_seq_field(paragraph, seq_name):
    """在段落中插入SEQ域"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f'SEQ {seq_name} \* ARABIC'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    return run


def add_caption_with_numbering(doc, title, is_table=False):
    """
    添加带自动编号的图表标题（SEQ域）
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    prefix = "表" if is_table else "图"
    run_prefix = para.add_run(prefix)
    run_prefix.font.bold = True
    run_prefix.font.size = Pt(10.5)
    run_prefix.font.name = 'Times New Roman'
    run_prefix._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 插入SEQ域
    add_seq_field(para, "Table" if is_table else "Figure")

    run_title = para.add_run(f" {title}")
    run_title.font.size = Pt(10.5)
    run_title.font.name = 'Times New Roman'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    return para


def add_image_caption_below(doc, title):
    """在图片下方添加标题（SEQ域）"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_prefix = para.add_run("图")
    run_prefix.font.bold = True
    run_prefix.font.size = Pt(10.5)
    run_prefix.font.name = 'Times New Roman'
    run_prefix._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 插入SEQ域
    add_seq_field(para, "Figure")

    run_title = para.add_run(f" {title}")
    run_title.font.size = Pt(10.5)
    run_title.font.name = 'Times New Roman'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    return para


def hide_table_borders(table):
    """隐藏表格边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')

    for border_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:val'), 'nil')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)

    tblPr.append(tblBorders)


def images_to_word(txt_path, output_word, images_folder="."):
    """
    将文本文件中的图片和Excel内容按指定格式插入到Word文档
    支持图表自动编号（SEQ域）
    """
    doc = Document()
    set_global_font(doc, font_name='宋体', font_size=Pt(10.5))

    with open(txt_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for group_num, line in enumerate(lines, 1):
        parts = [part.strip() for part in line.strip().split() if part.strip()]

        if len(parts) != 6:
            print(f"警告: 第{group_num}行不是6个元素，已跳过")
            continue

        image_names = parts[:5]
        excel_file = parts[5]

        first_image_serial = extract_serial_number(image_names[0])
        # group_heading = doc.add_heading(f"序号 {first_image_serial} 的数据与图表", level=2)
        # group_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # for run in group_heading.runs:
        #     run.font.name = 'Times New Roman'
        #     run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        #     run.font.size = Pt(14)
        #
        # doc.add_paragraph()

        def add_image_row(img_subset, description):
            serial_number = extract_serial_number(img_subset[0])
            title = f"校核点 {serial_number} {description}"

            table = doc.add_table(rows=1, cols=len(img_subset))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hide_table_borders(table)

            for i, img_name in enumerate(img_subset):
                img_path = os.path.join(images_folder, img_name)
                cell = table.cell(0, i)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                if not os.path.exists(img_path):
                    para = cell.add_paragraph(f"图片 {img_name} 不存在")
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    continue

                with Image.open(img_path) as img:
                    aspect_ratio = img.width / img.height

                target_height = Cm(5)
                target_width = Cm(5 * aspect_ratio)

                if len(img_subset) == 2 and target_width > Cm(7):
                    target_width = Cm(7)
                    target_height = Cm(target_width / aspect_ratio)

                para = cell.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(img_path, height=target_height, width=target_width)

            add_image_caption_below(doc, title)

        add_image_row(image_names[0:2], "的NAC图像（左）及其对应渲染DEM图像（右）")
        add_image_row(image_names[2:4], "的NAC图像（左）及其对应渲染DEM图像（右）二值化结果")
        add_image_row([image_names[4]], "光照失配图（绿色区域表示没有误差，黄色代表暗区失配，蓝色代表亮区失配）")

        excel_path = os.path.join(images_folder, excel_file)
        excel_serial = extract_serial_number(excel_file)
        table_title = f"校核点 {excel_serial} 的高程精度评价指标表"
        add_caption_with_numbering(doc, table_title, is_table=True)

        if not os.path.exists(excel_path):
            error_para = doc.add_paragraph(f"错误：Excel文件 '{excel_file}' 不存在")
            for run in error_para.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            wb = load_workbook(excel_path, data_only=True)
            ws = wb.active
            max_row, max_col = ws.max_row, ws.max_column

            table = doc.add_table(rows=max_row, cols=max_col)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'

            for row in range(max_row):
                for col in range(max_col):
                    cell_value = ws.cell(row=row + 1, column=col + 1).value
                    cell_text = str(cell_value) if cell_value is not None else ""
                    cell = table.cell(row, col)
                    cell.text = cell_text
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.size = Pt(12)
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            wb.close()

        # if group_num != len(lines):
        #     doc.add_paragraph()
        #     sep = doc.add_paragraph("-" * 50)
        #     sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #     for run in sep.runs:
        #         run.font.name = 'Times New Roman'
        #         run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        #     doc.add_paragraph()

    doc.save(output_word)
    print(f"已成功生成Word文档: {output_word}")
    print("提示：在Word中按 Ctrl+A → F9 可更新所有编号（图/表自动编号）")


if __name__ == "__main__":
    # txt_file = r"D:\长航程\校核用图\全部报告.txt"
    # output_file = r"D:\长航程\校核用图\结果汇总\结果汇总——09262.docx"
    # files_dir = r"D:\长航程\校核用图\失配结果汇总"
    txt_file = r"C:\Users\Lenovo\Desktop\新的校核\图像\生成汇总文档txt.txt"
    output_file = r"C:\Users\Lenovo\Desktop\新的校核\图像\502_0206.docx"
    files_dir = r"C:\Users\Lenovo\Desktop\新的校核\图像\全部汇总"

    images_to_word(txt_file, output_file, files_dir)
