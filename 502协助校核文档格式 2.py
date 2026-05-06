from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
from openpyxl import load_workbook
from PIL import Image


def set_global_font(doc, font_name='宋体', font_size=Pt(10.5)):
    """设置全局默认字体"""
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    doc.styles['Normal'].font.size = font_size


def extract_serial_number(filename):
    """从文件名中提取序号"""
    match = re.search(r'pt(\d+)', filename)
    if match:
        return match.group(1)
    else:
        return filename


def add_chapter_seq_field(paragraph, seq_name, chapter_num):
    """在段落中插入带章节号的SEQ域"""
    run = paragraph.add_run()

    # 添加章节号和连字符
    chapter_run = paragraph.add_run(f"{chapter_num}-")
    chapter_run.font.bold = True
    chapter_run.font.size = Pt(10.5)
    chapter_run.font.name = 'Times New Roman'
    chapter_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 添加SEQ域
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f'SEQ {seq_name} \* ARABIC'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    return run


def add_caption_with_numbering(doc, title, chapter_num, is_table=False):
    """添加带章节编号的图表标题（SEQ域）"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    prefix = "表" if is_table else "图"
    run_prefix = para.add_run(prefix)
    run_prefix.font.bold = True
    run_prefix.font.size = Pt(10.5)
    run_prefix.font.name = 'Times New Roman'
    run_prefix._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_chapter_seq_field(para, "Table" if is_table else "Figure", chapter_num)

    run_title = para.add_run(f" {title}")
    run_title.font.size = Pt(10.5)
    run_title.font.name = 'Times New Roman'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    return para


def reset_seq_numbering(doc, seq_name):
    """重置SEQ编号"""
    para = doc.add_paragraph()
    run = para.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f'SEQ {seq_name} \* ARABIC \r'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    pPr = para._element.get_or_add_pPr()
    vanish = OxmlElement('w:vanish')
    pPr.append(vanish)

    return para


def set_cell_text(cell, text, font_size=Pt(10.5), bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """设置单元格文本格式"""
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for para in cell.paragraphs:
        para.alignment = alignment
        for run in para.runs:
            run.font.size = font_size
            run.font.bold = bold
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_image_to_cell(cell, img_path, target_height=Cm(4.5)):
    """向单元格添加图片"""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    if not os.path.exists(img_path):
        para = cell.add_paragraph(f"图片不存在")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    with Image.open(img_path) as img:
        aspect_ratio = img.width / img.height

    target_width = Cm(target_height.cm * aspect_ratio)

    para = cell.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(img_path, height=target_height, width=target_width)


def images_to_word(txt_path, output_word, images_folder=".", chapter_num=1):
    """
    将文本文件中的图片和Excel内容按新格式插入到Word文档
    每行10个元素：8个图片文件名 + 2个Excel文件名
    图片排列：
    - 行1：NAC原图、河南渲染、香港渲染
    - 行2：NAC二值化、河南二值化、香港二值化
    - 行3：河南失配、香港失配
    """
    doc = Document()
    set_global_font(doc, font_name='宋体', font_size=Pt(10.5))

    # 添加章节标题
    chapter_heading = doc.add_heading(f"第{chapter_num}章 数据与图表", level=1)
    chapter_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in chapter_heading.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(16)
    doc.add_paragraph()

    # 重置本章的图表编号
    reset_seq_numbering(doc, "Figure")
    reset_seq_numbering(doc, "Table")

    with open(txt_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for group_num, line in enumerate(lines, 1):
        parts = [part.strip() for part in line.strip().split() if part.strip()]

        # 需要10个元素：8个图片 + 2个Excel文件
        if len(parts) != 10:
            print(f"警告: 第{group_num}行不是10个元素（8图片+2Excel），已跳过")
            continue

        image_names = parts[:8]  # 8张图片
        excel_files = parts[8:10]  # 2个Excel文件

        serial_number = extract_serial_number(image_names[0])

        # 创建图片大表格 (6行3列) - 不包含失配图
        # 第1行：总标题 "NAC 原图与渲染结果"
        # 第2行：列标题（NAC、河南大学渲染结果、香港理工渲染结果）
        # 第3行：3张原图
        # 第4行：二值化结果标题
        # 第5行：列标题
        # 第6行：3张二值化图

        main_table = doc.add_table(rows=6, cols=3)
        main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        main_table.style = 'Table Grid'

        # 第1行：总标题 "NAC 原图与渲染结果"
        cell_title = main_table.cell(0, 0)
        cell_title.merge(main_table.cell(0, 2))
        set_cell_text(cell_title, "NAC 原图与渲染结果", font_size=Pt(12), bold=True)

        # 第2行：列标题
        set_cell_text(main_table.cell(1, 0), "NAC", bold=True)
        set_cell_text(main_table.cell(1, 1), "河南大学渲染结果", bold=True)
        set_cell_text(main_table.cell(1, 2), "香港理工渲染结果", bold=True)

        # 第3行：前3张图片（索引0, 1, 2）
        for i in range(3):
            img_path = os.path.join(images_folder, image_names[i])
            add_image_to_cell(main_table.cell(2, i), img_path)

        # 第4行：二值化结果标题
        cell_binary_title = main_table.cell(3, 0)
        cell_binary_title.merge(main_table.cell(3, 2))
        set_cell_text(cell_binary_title, "二值化结果图", font_size=Pt(12), bold=True)

        # 第5行：列标题
        set_cell_text(main_table.cell(4, 0), "NAC", bold=True)
        set_cell_text(main_table.cell(4, 1), "河南大学", bold=True)
        set_cell_text(main_table.cell(4, 2), "香港理工大学", bold=True)

        # 第6行：3张二值化图片（索引3, 4, 5）
        for i in range(3):
            img_path = os.path.join(images_folder, image_names[3 + i])
            add_image_to_cell(main_table.cell(5, i), img_path)

        # 创建失配图表格（独立的2列表格）
        mismatch_table = doc.add_table(rows=3, cols=2)
        mismatch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        mismatch_table.style = 'Table Grid'

        # 第1行：失配图结果标题（合并2列）
        cell_mismatch_title = mismatch_table.cell(0, 0)
        cell_mismatch_title.merge(mismatch_table.cell(0, 1))
        set_cell_text(cell_mismatch_title, "失配图结果", font_size=Pt(12), bold=True)

        # 第2行：列标题
        set_cell_text(mismatch_table.cell(1, 0), "河南大学", bold=True)
        set_cell_text(mismatch_table.cell(1, 1), "香港理工大学", bold=True)

        # 第3行：2张失配图（索引6, 7）
        for i in range(2):
            img_path = os.path.join(images_folder, image_names[6 + i])
            add_image_to_cell(mismatch_table.cell(2, i), img_path)

        # 添加图片标题
        title = f"校核点 {serial_number} 高程校核图"
        add_caption_with_numbering(doc, title, chapter_num, is_table=False)
        doc.add_paragraph()

        # 创建合并的Excel表格
        excel_data = []
        for excel_file in excel_files:
            excel_path = os.path.join(images_folder, excel_file)
            if os.path.exists(excel_path):
                wb = load_workbook(excel_path, data_only=True)
                ws = wb.active
                data = []
                for row in ws.iter_rows(values_only=True):
                    data.append(list(row))
                excel_data.append(data)
                wb.close()
            else:
                print(f"警告: Excel文件 '{excel_file}' 不存在")
                excel_data.append([])

        if len(excel_data) == 2 and excel_data[0] and excel_data[1]:
            # 确保两个表格行数相同
            max_rows = max(len(excel_data[0]), len(excel_data[1]))

            # 创建合并表格：第一列共享 + 两组结果列
            cols_1 = len(excel_data[0][0]) if excel_data[0] else 0
            cols_2 = len(excel_data[1][0]) if excel_data[1] else 0
            total_cols = 1 + (cols_1 - 1) + (cols_2 - 1)  # 共享第一列

            # 添加表格标题
            table_title = f"校核点 {serial_number} 的高程精度评价指标表"
            add_caption_with_numbering(doc, table_title, chapter_num, is_table=True)

            merged_table = doc.add_table(rows=max_rows, cols=total_cols)
            merged_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            merged_table.style = 'Table Grid'

            # 填充数据
            for row_idx in range(max_rows):
                col_offset = 0

                # 第一列（计算指标）- 从第一个Excel取
                if row_idx < len(excel_data[0]) and len(excel_data[0][row_idx]) > 0:
                    cell_value = excel_data[0][row_idx][0]
                    set_cell_text(merged_table.cell(row_idx, col_offset),
                                  str(cell_value) if cell_value is not None else "",
                                  font_size=Pt(10))
                col_offset += 1

                # 第一个Excel的其余列
                if row_idx < len(excel_data[0]):
                    for col_idx in range(1, min(cols_1, len(excel_data[0][row_idx]))):
                        cell_value = excel_data[0][row_idx][col_idx]
                        set_cell_text(merged_table.cell(row_idx, col_offset),
                                      str(cell_value) if cell_value is not None else "",
                                      font_size=Pt(10))
                        col_offset += 1
                else:
                    col_offset += (cols_1 - 1)

                # 第二个Excel的数据列（跳过第一列）
                if row_idx < len(excel_data[1]):
                    for col_idx in range(1, min(cols_2, len(excel_data[1][row_idx]))):
                        cell_value = excel_data[1][row_idx][col_idx]
                        set_cell_text(merged_table.cell(row_idx, col_offset),
                                      str(cell_value) if cell_value is not None else "",
                                      font_size=Pt(10))
                        col_offset += 1

        # 添加分隔符
        # if group_num != len(lines):
        #     doc.add_paragraph()
        #     sep = doc.add_paragraph("-" * 50)
        #     sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #     for run in sep.runs:
        #         run.font.name = 'Times New Roman'
        #         run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        #     doc.add_paragraph()

    doc.save(output_word)
    print(f"已成功生成Word文档: {output_word}")
    print("提示：在Word中按 Ctrl+A → F9 可更新所有编号（图/表自动编号）")


if __name__ == "__main__":
    txt_file = r"D:\桌面\sb\河南大学结果汇总\tif——all.txt"
    output_file = r"D:\桌面\sb\河南大学结果汇总\24日_new2.docx"
    files_dir = r"D:\桌面\sb\河南大学结果汇总"

    # TXT文件格式：每行10个元素
    # 图片顺序（8张）：
    # [0]NAC原图 [1]河南渲染 [2]香港渲染 [3]NAC二值化 [4]河南二值化 [5]香港二值化 [6]河南失配 [7]香港失配
    # Excel文件（2个）：
    # [8]河南Excel [9]香港Excel
    images_to_word(txt_file, output_file, files_dir, chapter_num=1)